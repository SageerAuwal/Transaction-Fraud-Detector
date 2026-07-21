import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page Margins (1 inch everywhere)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Fonts
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Color Palette
    PRIMARY_NAVY = RGBColor(0x0F, 0x17, 0x2A)  # #0f172a
    ACCENT_CYAN  = RGBColor(0x02, 0x84, 0xC7)  # #0284c7
    PURPLE_ACCENT= RGBColor(0x7C, 0x3A, 0xED)  # #7c3aed
    DARK_TEXT    = RGBColor(0x1E, 0x29, 0x3B)  # #1e293b
    MUTED_GRAY   = RGBColor(0x64, 0x74, 0x8B)  # #64748b

    # Title Banner / Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("GOJO SENTINEL")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(30)
    run_title.font.bold = True
    run_title.font.color.rgb = ACCENT_CYAN

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Comprehensive System Architecture, Research Design, Implementation & Operational Guide")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = MUTED_GRAY

    doc.add_paragraph() # spacing

    # Executive Metadata Box Table
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("System Name:", "Gojo Sentinel Transaction Fraud Detector"),
        ("Target Industry:", "Nigerian Banking Sector (CBN / NIBSS / NIP / USSD Compliance)"),
        ("Core Architecture:", "Hybrid Dual-Engine: Statutory Rules Engine + XGBoost AI Machine Learning"),
        ("Supported Platforms:", "Web Desktop Application & Apache Cordova Android Mobile APK")
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        cell_k = table_meta.cell(row_idx, 0)
        cell_v = table_meta.cell(row_idx, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, "F1F5F9")
        set_cell_background(cell_v, "F8FAFC")
        
        pk = cell_k.paragraphs[0]
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.color.rgb = PRIMARY_NAVY
        
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.color.rgb = DARK_TEXT

    doc.add_paragraph() # spacing

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_NAVY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = ACCENT_CYAN
        return h

    def add_body_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.color.rgb = DARK_TEXT
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        run_b = p.add_run(bold_prefix + ": ")
        run_b.bold = True
        run_b.font.color.rgb = PRIMARY_NAVY
        run_t = p.add_run(text)
        run_t.font.color.rgb = DARK_TEXT

    # 1. PROBLEM STATEMENT OF THE STUDY
    add_heading_1("1. Problem Statement of the Study")
    add_body_p(
        "The rapid digitization of the Nigerian financial services sector—driven by the Central Bank of Nigeria (CBN) Cashless Policy and NIBSS Instant Payments (NIP)—has led to exponential growth in electronic transaction volumes. However, this expansion has been accompanied by a sophisticated rise in electronic financial fraud, costing Nigerian financial institutions billions of Naira annually."
    )
    add_body_p(
        "Traditional fraud detection mechanisms in Nigerian banks rely primarily on static rule-based systems. While effective at enforcing hard statutory limits (such as daily transfer caps), static rules fail to adapt to complex, evolving fraud tactics like SIM-swap scams, credential harvesting, velocity attacks, and midnight account draining. Conversely, pure machine learning approaches often operate as 'black boxes' that ignore statutory Central Bank of Nigeria (CBN) regulatory thresholds, leading to compliance risk and non-compliance fines."
    )
    add_body_p(
        "Therefore, there is a critical need for a unified, hybrid fraud prevention system tailored to the Nigerian financial context that simultaneously guarantees 100% regulatory compliance while dynamically detecting complex fraud patterns with machine learning."
    )

    # 2. AIM AND OBJECTIVES OF THE STUDY
    add_heading_1("2. Aim & Objectives of the Study")
    add_body_p(
        "The primary aim of this project is to design, develop, and deploy 'Gojo Sentinel'—a hybrid transaction fraud detection and regulatory compliance platform specifically engineered for Nigerian banking channels."
    )
    add_body_p("To achieve this aim, the study pursued five specific objectives:")
    add_bullet("Objective 1 — Enforce Regulatory Compliance", "Implement a deterministic Nigerian Banking Compliance Engine enforcing CBN operational caps (USSD ₦100,000 daily limit, NIP ₦5,000,000 single transfer cap, midnight velocity checks, and restricted bank watchlists).")
    add_bullet("Objective 2 — Train Machine Learning Model", "Develop and train an XGBoost (Extreme Gradient Boosting) classifier to evaluate non-linear feature interactions and compute probabilistic fraud risk scores.")
    add_bullet("Objective 3 — Dual-Mode Screening Pipeline", "Build a unified pipeline supporting both real-time single transaction scoring and high-volume CSV dataset batch screening with live audit terminal logs.")
    add_bullet("Objective 4 — Modern Multi-Device Interface", "Design a responsive frontend UI featuring a Gemini-style Collapsible Rail Sidebar, 30-Day Activity Trend Analytics, and an Apache Cordova Android mobile APK package.")
    add_bullet("Objective 5 — Enterprise Security & Access Control", "Implement role-based security using JWT authentication, admin rule management, and analyst monitoring privileges.")

    # 3. WHAT THE SYSTEM DOES
    add_heading_1("3. What the System Does")
    add_body_p(
        "Gojo Sentinel acts as an intelligent intermediary between transaction channels and financial ledger backends. It screens every transaction against statutory rules and machine learning models to generate sub-second risk verdicts."
    )

    table_verdicts = doc.add_table(rows=5, cols=4)
    table_verdicts.alignment = WD_TABLE_ALIGNMENT.CENTER
    v_headers = ["Risk Level", "Probability Range", "Verdict", "Operational System Behavior"]
    for i, h_text in enumerate(v_headers):
        cell = table_verdicts.cell(0, i)
        set_cell_background(cell, "0F172A")
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    v_data = [
        ("LOW RISK", "0.0% - 29.9%", "APPROVE", "Transaction is processed instantly without user friction."),
        ("MEDIUM RISK", "30.0% - 59.9%", "REVIEW", "Held for step-up multi-factor authentication (2FA / OTP check)."),
        ("HIGH RISK", "60.0% - 84.9%", "DECLINE", "Transaction declined; suspicious activity alert logged."),
        ("CRITICAL RISK", "85.0% - 100.0%", "BLOCK", "Transaction blocked immediately due to rule violation or critical fraud score.")
    ]
    for row_idx, data in enumerate(v_data, start=1):
        bg = "F0FDF4" if row_idx == 1 else "FEFCE8" if row_idx == 2 else "FFF7ED" if row_idx == 3 else "FEF2F2"
        for col_idx, text in enumerate(data):
            cell = table_verdicts.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            p.add_run(text)

    doc.add_paragraph()

    # 4. FRONTEND ARCHITECTURE & TOOLS USED
    add_heading_1("4. Frontend Architecture & Technology Stack")
    add_body_p(
        "The frontend is engineered as an ultra-fast, responsive Single Page Application (SPA) built with vanilla core web technologies for maximum performance and cross-platform compatibility."
    )

    add_bullet("Core Technologies", "HTML5 for semantic layout, Vanilla CSS3 with custom HSL tokenized design system (dark mode, glassmorphism), and Vanilla JavaScript (ES6+ async/await, Fetch API).")
    add_bullet("Gemini-Style Collapsible Rail Sidebar", "Supports Expanded Mode (260px with logo branding and labels) and Collapsible Rail Mode (72px compact icon bar with centered SVGs and tooltips).")
    add_bullet("Real-Time Terminal Console & Progress Bar", "Includes an animated 0-100% progress bar alongside a dark green live terminal logger outputting timestamped execution events during dataset batch processing.")
    add_bullet("30-Day Activity Trend & Doughnut Analytics", "Integrated Chart.js library rendering real-time line charts for legitimate vs. fraudulent volume trends and interactive risk doughnut breakdown charts.")
    add_bullet("Vector SVG Graphic System", "100% vector SVG icons across all navigation, step visualizer cards, and violation badges for crisp display across mobile and desktop displays.")

    # 5. BACKEND ARCHITECTURE & TOOLS USED
    add_heading_1("5. Backend Architecture & Technology Stack")
    add_body_p(
        "The backend is powered by a high-performance Python microservice architecture built on FastAPI and Uvicorn:"
    )

    add_bullet("FastAPI & Uvicorn", "Asynchronous Python web framework delivering high concurrency and automatic OpenAPI / Swagger documentation.")
    add_bullet("XGBoost Classifier (`xgboost`)", "Extreme Gradient Boosting algorithm trained on high-dimensional transaction features to compute exact probability scores.")
    add_bullet("Data Processing (`pandas`, `numpy`, `scikit-learn`)", "Data manipulation and feature preprocessing pipeline for real-time payloads and bulk CSV files.")
    add_bullet("Authentication & Security (`PyJWT`, `passlib`)", "JSON Web Token (JWT) stateless authentication and bcrypt password hashing for admin security.")
    add_bullet("CORS Middleware", "Configured cross-origin resource sharing allowing secure API requests from web browsers and mobile apps.")

    # 6. STEP-BY-STEP OPERATIONAL GUIDE
    add_heading_1("6. Detailed Operational & User Guide")
    add_body_p("Follow this step-by-step operational guide to run and use Gojo Sentinel:")

    add_heading_2("6.1 How to Start the System Locally")
    add_bullet("Windows Execution", "Double-click `run_offline.bat` in the project root directory. It automatically starts the Uvicorn FastAPI backend on port 8000 and opens the web app in your browser.")
    add_bullet("Linux / macOS Execution", "Open terminal and execute `./run_linux.sh`.")

    add_heading_2("6.2 Scoring Single Transactions (Simulator)")
    add_bullet("Step 1", "Navigate to the 'Dashboard' section in the sidebar.")
    add_bullet("Step 2", "Select a channel (NIP, USSD, POS, Mobile, Web).")
    add_bullet("Step 3", "Click 'Normal Pattern' to fill safe parameters, or 'High-Risk Pattern' to simulate an attack (e.g. USSD ₦250,000 at 02:15 AM).")
    add_bullet("Step 4", "Click 'Score Transaction' to view immediate AI Risk %, Decision Badge, and Compliance Violation tags.")

    add_heading_2("6.3 Bulk Dataset Batch Screening")
    add_bullet("Step 1", "Click 'Batch Scanner' in the sidebar.")
    add_bullet("Step 2", "Drag & drop a transaction dataset CSV file (or click 'Select CSV File').")
    add_bullet("Step 3", "Observe the live 4-Step Visualizer Boxes light up in glowing cyan, gold, purple, and green as the Live Terminal Console logs processing timestamps.")
    add_bullet("Step 4", "Review the 4 Stat Cards, Risk Doughnut Chart, Top Rule Breaches, and Filterable Audited Table.")
    add_bullet("Step 5", "Click 'Export Audited CSV' to download the completed audit report.")

    add_heading_2("6.4 Managing Rules & Users (Admin)")
    add_bullet("Step 1", "Click 'Admin Access' on the sidebar user card and log in with admin credentials (`admin` / `admin123`).")
    add_bullet("Step 2", "Navigate to 'Rules' to add new USSD/NIP thresholds or delete rules.")
    add_bullet("Step 3", "Navigate to 'Users' to register new analyst accounts.")

    add_heading_2("6.5 Mobile Android APK")
    add_bullet("Execution", "Open `mobile/platforms/android/app/build/outputs/apk/debug/app-debug.apk` to install directly on Android mobile devices.")

    # Save document
    filename_docx = "Gojo_Sentinel_Complete_System_Presentation.docx"
    doc.save(filename_docx)
    print(f"[SUCCESS] Microsoft Word document created successfully: {filename_docx}")

    # Also generate comprehensive Markdown version
    md_content = f"""# 🛡️ GOJO SENTINEL
## Comprehensive System Architecture, Research Design, Implementation & Operational Guide

---

### **1. Problem Statement of the Study**
The rapid digitization of the Nigerian financial services sector—driven by the Central Bank of Nigeria (CBN) Cashless Policy and NIBSS Instant Payments (NIP)—has led to exponential growth in electronic transaction volumes. However, this expansion has been accompanied by a sophisticated rise in electronic financial fraud, costing Nigerian financial institutions billions of Naira annually.

Traditional fraud detection mechanisms in Nigerian banks rely primarily on static rule-based systems. While effective at enforcing hard statutory limits (such as daily transfer caps), static rules fail to adapt to complex, evolving fraud tactics like SIM-swap scams, credential harvesting, velocity attacks, and midnight account draining. Conversely, pure machine learning approaches often operate as "black boxes" that ignore statutory Central Bank of Nigeria (CBN) regulatory thresholds, leading to compliance risk and non-compliance fines.

Therefore, there is a critical need for a unified, hybrid fraud prevention system tailored to the Nigerian financial context that simultaneously guarantees 100% regulatory compliance while dynamically detecting complex fraud patterns with machine learning.

---

### **2. Aim & Objectives of the Study**
The primary aim of this project is to design, develop, and deploy **Gojo Sentinel**—a hybrid transaction fraud detection and regulatory compliance platform specifically engineered for Nigerian banking channels.

- **Objective 1 — Enforce Regulatory Compliance**: Implement a deterministic Nigerian Banking Compliance Engine enforcing CBN operational caps (USSD ₦100,000 daily limit, NIP ₦5,000,000 single transfer cap, midnight velocity checks, and restricted bank watchlists).
- **Objective 2 — Train Machine Learning Model**: Develop and train an XGBoost (Extreme Gradient Boosting) classifier to evaluate non-linear feature interactions and compute probabilistic fraud risk scores.
- **Objective 3 — Dual-Mode Screening Pipeline**: Build a unified pipeline supporting both real-time single transaction scoring and high-volume CSV dataset batch screening with live audit terminal logs.
- **Objective 4 — Modern Multi-Device Interface**: Design a responsive frontend UI featuring a Gemini-style Collapsible Rail Sidebar, 30-Day Activity Trend Analytics, and an Apache Cordova Android mobile APK package.
- **Objective 5 — Enterprise Security & Access Control**: Implement role-based security using JWT authentication, admin rule management, and analyst monitoring privileges.

---

### **3. What the System Does (Risk Verdict Matrix)**

| Risk Level | Fraud Probability Range | Action Verdict | Operational System Behavior |
| :--- | :--- | :--- | :--- |
| **LOW RISK** | 0.0% - 29.9% | **`APPROVE`** | Transaction is processed instantly without user friction. |
| **MEDIUM RISK** | 30.0% - 59.9% | **`REVIEW`** | Held for step-up multi-factor authentication (2FA / OTP check). |
| **HIGH RISK** | 60.0% - 84.9% | **`DECLINE`** | Transaction declined; suspicious activity alert logged. |
| **CRITICAL RISK** | 85.0% - 100.0% | **`BLOCK`** | Transaction blocked immediately due to rule violation or critical fraud score. |

---

### **4. Frontend Architecture & Technology Stack**
- **Core Web Stack**: HTML5, Vanilla CSS3 (HSL dark mode design system, glassmorphism), Vanilla JavaScript (ES6+ async/await, Fetch API).
- **Gemini-Style Collapsible Rail Sidebar**: Supports Expanded Mode (260px with branding & text labels) and Collapsible Rail Mode (72px compact icon bar with centered SVGs and tooltips).
- **Real-Time Terminal Console & Progress Bar**: Animated 0-100% progress bar alongside a dark green live terminal logger outputting timestamped execution events during dataset batch processing.
- **Analytics & Visualizations**: Chart.js rendering 30-Day Activity Trends and interactive Risk Distribution Doughnut charts.
- **Vector SVG System**: 100% scalable SVG vector graphics across all icons, cards, and badges.

---

### **5. Backend Architecture & Technology Stack**
- **FastAPI & Uvicorn**: Asynchronous Python web framework delivering high concurrency.
- **XGBoost Classifier (`xgboost`)**: Extreme Gradient Boosting algorithm computing exact probability scores.
- **Data Preprocessing (`pandas`, `numpy`, `scikit-learn`)**: Data pipeline for real-time and bulk CSV records.
- **Security (`PyJWT`, `passlib`)**: JWT stateless authentication and bcrypt password hashing.

---

### **6. Detailed Step-by-Step User Guide**

#### **6.1 How to Start the System Locally**
- **Windows**: Double-click `run_offline.bat` in the project folder.
- **Linux / macOS**: Execute `./run_linux.sh` in terminal.

#### **6.2 Scoring Single Transactions**
1. Open the **Dashboard** tab.
2. Select a channel (NIP, USSD, POS, Mobile, Web).
3. Click **Normal Pattern** or **High-Risk Pattern** to fill sample data.
4. Click **Score Transaction** to view immediate Risk %, Decision Badge, and Compliance Violation tags.

#### **6.3 Bulk Dataset Batch Screening**
1. Open the **Batch Scanner** tab.
2. Drag & drop a dataset CSV file.
3. Observe the 4-Step Visualizer Boxes glow in real-time as the Live Terminal Console streams execution timestamps.
4. View the 4 Summary Stat Cards, Doughnut Chart, Top Rule Breaches, and Filterable Table.
5. Click **Export Audited CSV** to download the completed audit report.

#### **6.4 Admin Management**
1. Click **Admin Access** on the user card and log in (`admin` / `admin123`).
2. Manage rules under **Rules** tab or add users under **Users** tab.
"""

    with open("GOJO_SENTINEL_SYSTEM_DOCUMENTATION.md", "w", encoding="utf-8") as f:
        f.write(md_content)
    print("[SUCCESS] Markdown documentation created successfully: GOJO_SENTINEL_SYSTEM_DOCUMENTATION.md")

if __name__ == "__main__":
    create_document()
